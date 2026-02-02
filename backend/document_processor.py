"""
Document Processing Module for AgreeWise
Handles text extraction from DOCX, PDF, and images using OCR
"""

import os
import magic
from docx import Document
import pdfplumber
from pdf2image import convert_from_path
import easyocr
from PIL import Image
import tempfile

# Initialize EasyOCR reader (lazy load to avoid startup delay)
_ocr_reader = None
_ocr_reader_languages = None


def get_ocr_reader(languages=['en']):
    """Get or initialize EasyOCR reader with specified languages"""
    global _ocr_reader, _ocr_reader_languages

    # Reinitialize if language changes
    if _ocr_reader is None or _ocr_reader_languages != languages:
        print(f'🔧 Initializing EasyOCR with languages: {languages}')
        _ocr_reader = easyocr.Reader(languages, gpu=False)  # Use CPU mode
        _ocr_reader_languages = languages
        print(f'✓ EasyOCR ready for {languages}')

    return _ocr_reader


def detect_file_type(file_path):
    """Detect file type using python-magic"""
    mime_type = magic.from_file(file_path, mime=True)
    return mime_type


def extract_from_docx(file_path):
    """
    Extract text from DOCX files
    Includes paragraphs, tables, headers, and footers
    """
    print('📄 Extracting text from DOCX...')
    doc = Document(file_path)

    full_text = []

    # Extract all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Extract tables (important for contract terms!)
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    extracted_text = '\n\n'.join(full_text)
    print(f'✓ Extracted {len(extracted_text)} characters from DOCX')
    return extracted_text


def extract_from_pdf(file_path):
    """
    Extract text from text-based PDFs using pdfplumber
    Returns empty string if PDF is scanned (no extractable text)
    """
    print('📄 Extracting text from PDF...')

    full_text = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    full_text.append(text)
                    print(f'  Page {page_num}: {len(text)} characters')

    except Exception as e:
        print(f'❌ PDF extraction error: {str(e)}')
        return ''

    extracted_text = '\n\n'.join(full_text)

    # Check if we got meaningful text (scanned PDFs might return garbage)
    if len(extracted_text.strip()) < 50:
        print('⚠️  Very little text extracted - PDF might be scanned')
        return ''

    print(f'✓ Extracted {len(extracted_text)} characters from PDF')
    return extracted_text


def ocr_image(file_path, language='en'):
    """
    Perform OCR on an image file using EasyOCR
    """
    print(f'🔍 Running OCR on image (language: {language})...')

    # Map common language codes to EasyOCR codes
    lang_map = {
        'en': 'en', 'es': 'es', 'fr': 'fr', 'de': 'de', 'pt': 'pt',
        'zh': 'ch_sim', 'ja': 'ja', 'ko': 'ko', 'ar': 'ar', 'hi': 'hi',
        'ru': 'ru', 'it': 'it', 'tr': 'tr', 'pl': 'pl', 'nl': 'nl'
    }

    ocr_lang = lang_map.get(language, 'en')

    try:
        # Get OCR reader
        reader = get_ocr_reader([ocr_lang])

        # Perform OCR
        results = reader.readtext(file_path, detail=0, paragraph=True)

        # Join all detected text blocks
        extracted_text = '\n\n'.join(results)

        print(f'✓ OCR completed: {len(extracted_text)} characters extracted')
        return extracted_text

    except Exception as e:
        print(f'❌ OCR error: {str(e)}')
        raise


def ocr_pdf(file_path, language='en'):
    """
    Convert PDF pages to images and perform OCR
    Used for scanned PDFs with no extractable text
    """
    print(f'🔍 Converting PDF to images for OCR (language: {language})...')

    try:
        # Convert PDF to images
        images = convert_from_path(file_path, dpi=300)  # High DPI for better OCR
        print(f'  Converted to {len(images)} images')

        full_text = []

        # OCR each page
        for page_num, image in enumerate(images, 1):
            print(f'  OCR Page {page_num}...')

            # Save image temporarily
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                image.save(temp_img.name, 'PNG')
                temp_path = temp_img.name

            try:
                # Perform OCR
                page_text = ocr_image(temp_path, language)
                if page_text:
                    full_text.append(page_text)
            finally:
                # Clean up temp image
                os.remove(temp_path)

        extracted_text = '\n\n'.join(full_text)
        print(f'✓ PDF OCR completed: {len(extracted_text)} characters')
        return extracted_text

    except Exception as e:
        print(f'❌ PDF OCR error: {str(e)}')
        raise


def clean_text(text):
    """
    Clean and normalize extracted text
    Remove excessive whitespace, normalize line breaks
    """
    if not text:
        return ''

    # Replace multiple spaces with single space
    text = ' '.join(text.split())

    # Normalize line breaks (keep paragraph structure)
    lines = text.split('\n')
    lines = [line.strip() for line in lines if line.strip()]

    return '\n\n'.join(lines)


def process_document(file_path, language='en'):
    """
    Main document processing function
    Detects file type and extracts text using appropriate method

    Args:
        file_path: Path to the document file
        language: Language code for OCR (default: 'en')

    Returns:
        dict: {
            'text': extracted text,
            'method': extraction method used,
            'file_type': detected MIME type,
            'char_count': number of characters,
            'success': True/False
        }
    """

    try:
        # Detect file type
        mime_type = detect_file_type(file_path)
        print(f'📁 Detected file type: {mime_type}')

        extracted_text = ''
        method = ''

        # DOCX - Direct text extraction
        if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            extracted_text = extract_from_docx(file_path)
            method = 'docx_extraction'

        # PDF - Try text extraction, fallback to OCR
        elif mime_type == 'application/pdf':
            extracted_text = extract_from_pdf(file_path)

            if not extracted_text or len(extracted_text.strip()) < 50:
                print('📸 PDF appears to be scanned, using OCR...')
                extracted_text = ocr_pdf(file_path, language)
                method = 'pdf_ocr'
            else:
                method = 'pdf_extraction'

        # Images - OCR required
        elif mime_type.startswith('image/'):
            extracted_text = ocr_image(file_path, language)
            method = 'image_ocr'

        else:
            raise ValueError(f'Unsupported file type: {mime_type}')

        # Clean the extracted text
        cleaned_text = clean_text(extracted_text)

        return {
            'success': True,
            'text': cleaned_text,
            'method': method,
            'file_type': mime_type,
            'char_count': len(cleaned_text),
            'error': None
        }

    except Exception as e:
        print(f'❌ Document processing error: {str(e)}')
        return {
            'success': False,
            'text': '',
            'method': None,
            'file_type': None,
            'char_count': 0,
            'error': str(e)
        }
